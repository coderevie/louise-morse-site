"""Build a Word document of the readings that need a second pair of eyes.

For each word the checker cannot settle on its own, this cuts the strip of
the scan the word sits on, prints what the text currently says and what the
alternative would be, and leaves a line to write what the page actually
shows.

The scans carry no text layer, so a word's place on the page is worked out
from how far down the OCR's own lines it falls, measured against the band of
the page that actually holds ink. That lands within a line or two, and the
strip is cut generously enough to take that in.

    python tools/make_review_doc.py [first] [count]

Writes review/wastena-review-<first>-<last>.docx and the strips beside it.
"""
import os
import re
import sys

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import proofread_wastena as P  # noqa: E402

SCAN = os.path.join("pdfs", "wastena", "wastena-full.pdf")
SOURCE = os.path.join("transcripts", "chatgpt-source")
OUT = "review"
WIDTH = 1250          # pixels across the cut strip
QUALITY = 70          # the strips are grey typing; JPEG carries it well
LINES_EITHER_SIDE = 5


def raw_pages():
    """The scanner's own output, page by page, with its line breaks kept."""
    pages, n = {}, 0
    for i in range(1, 11):
        path = os.path.join(SOURCE, f"the-wastena-retreat_p{i}.txt")
        text = open(path, encoding="utf-8", errors="replace").read()
        parts = re.split(r"===== PAGE (\d+) =====", text)
        for j in range(1, len(parts), 2):
            n += 1
            pages[n] = parts[j + 1]
    return pages


def ink_band(page):
    """The top and bottom of the part of the page that carries typing.

    Only the middle of the sheet is measured. These are photographs of loose
    pages, so the edges carry the shadow of the desk and the dark of the
    binder holes, and taking those in would make every page look inked from
    top to bottom.
    """
    pix = page.get_pixmap(colorspace=pymupdf.csGRAY, dpi=40)
    x0, x1 = int(pix.width * 0.18), int(pix.width * 0.82)
    y0, y1 = int(pix.height * 0.04), int(pix.height * 0.96)
    rows = []
    for y in range(y0, y1):
        row = [pix.pixel(x, y)[0] for x in range(x0, x1, 3)]
        rows.append((y, sum(row) / len(row)))
    if len(rows) < 10:
        return 0.10, 0.82

    values = sorted(v for _, v in rows)
    paper = values[int(len(values) * 0.9)]        # the bare sheet
    inked = [y for y, v in rows if v < paper - 12]
    if len(inked) < 8:
        return 0.10, 0.82
    return inked[0] / pix.height, (inked[-1] + 1) / pix.height


def strip_for(doc, page_no, line_no, line_count, out_path,
              lines_either_side=None):
    """Cut the strip of the scan that holds a given line of the page."""
    page = doc[page_no - 1]
    top, bottom = ink_band(page)
    height = bottom - top
    # Where this line falls down the inked band, plus a little either side.
    centre = top + height * ((line_no + 0.5) / max(line_count, 1))
    either = LINES_EITHER_SIDE if lines_either_side is None else lines_either_side
    span = height * (either * 2 + 1) / max(line_count, 1)
    span = min(max(span, 0.10), 0.34)
    y0 = max(0.0, centre - span / 2)
    y1 = min(1.0, centre + span / 2)

    rect = page.rect
    clip = pymupdf.Rect(rect.x0, rect.y0 + rect.height * y0,
                        rect.x1, rect.y0 + rect.height * y1)
    scale = WIDTH / rect.width
    pix = page.get_pixmap(matrix=pymupdf.Matrix(scale, scale), clip=clip,
                          colorspace=pymupdf.csGRAY)
    with open(out_path, "wb") as f:
        f.write(pix.tobytes("jpeg", jpg_quality=QUALITY))
    return out_path


def find_word(pages, word):
    """Where a word sits: its page, its line, and how many lines that page has.

    Only lines carrying something are counted, since the blank ones the
    scanner leaves between paragraphs have no height on the page and would
    otherwise push the estimate down.
    """
    pattern = re.compile(r"\b" + re.escape(word) + r"\b", re.I)
    for page_no, text in pages.items():
        lines = [ln for ln in text.splitlines() if ln.strip()]
        for i, line in enumerate(lines):
            if pattern.search(line):
                return page_no, i, len(lines), line.strip()
    return None


def sentence_for(word, texts):
    """The word in its sentence, as the site currently prints it."""
    for text in texts.values():
        m = re.search(r"\b" + re.escape(word) + r"\b", text, re.I)
        if m:
            start = max(0, m.start() - 130)
            end = min(len(text), m.end() + 130)
            return re.sub(r"\s+", " ", text[start:end]).strip()
    return ""


def shade(run, rgb):
    run.font.color.rgb = RGBColor(*rgb)


def garbled_blocks(texts, words):
    """Passages the scanner did not resolve into readable prose.

    Only the discourses are looked at. The session jottings on pages 7 to 41
    are notes rather than sentences, and the foreword on 42 to 45 carries
    place names and publishing terms, so both score badly without being
    damaged at all.
    """
    good = {w for w, c in words.items() if c >= 6}
    out = []
    for path, text in sorted(texts.items()):
        page = None
        for block in text.split("\n\n"):
            b = block.strip()
            mark = re.fullmatch(r"\[page (\d+)\]", b)
            if mark:
                page = int(mark.group(1))
                continue
            if not b or b.startswith("#") or page is None or page < 46:
                continue
            toks = re.findall(r"\b[a-z']{2,}\b", b)
            if len(toks) < 8:
                continue
            bad = [w for w in toks if w not in good]
            if len(bad) / len(toks) >= 0.20:
                out.append((page, b, len(bad) / len(toks)))
    out.sort(key=lambda r: -r[2])
    return out


def block_strip(scan, pages, page_no, block, out_path):
    """Cut the strip of scan holding a passage, found by its opening words."""
    opening = re.findall(r"[A-Za-z]{3,}", block)[:6]
    lines = [ln for ln in pages[page_no].splitlines() if ln.strip()]
    best, score = 0, -1
    for i, line in enumerate(lines):
        hit = sum(1 for w in opening if w.lower() in line.lower())
        if hit > score:
            best, score = i, hit
    span = max(len(re.findall(r"[A-Za-z]{3,}", block)) // 9, 4)
    return strip_for(scan, page_no, best + span // 2, len(lines), out_path,
                     lines_either_side=span)


def main():
    count = int(sys.argv[1]) if len(sys.argv) > 1 else 45

    texts = P.load()
    words = P.vocabulary(texts)
    suspects = [r for r in P.remaining(texts, words) if r[3]]
    suspects.sort(key=lambda r: (-r[1], r[0]))
    batch = suspects[:count]
    damaged = garbled_blocks(texts, words)

    os.makedirs(OUT, exist_ok=True)
    pages = raw_pages()
    scan = pymupdf.open(SCAN)

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)

    doc.add_heading("The Wastena Retreat — readings to check", 0)
    intro = doc.add_paragraph()
    intro.add_run(
        "Two kinds of question, in two parts. Part one is the handful of "
        "passages the scanner could not read at all, where a retyping would "
        "recover text that is currently lost. Part two is single words where "
        "the scan could be read either way. In both, please write what you "
        "see on the line marked "
    )
    intro.add_run("On the page:").bold = True
    intro.add_run(
        ". Where what the site prints is already right, \"ok\" is enough."
    )

    # ---- part one: passages the scanner could not resolve -----------------
    doc.add_heading("Part one — passages to retype", 1)
    doc.add_paragraph(
        f"{len(damaged)} passages, on pages "
        f"{', '.join(str(p) for p in sorted({d[0] for d in damaged}))}. "
        "These are the only places in the whole retreat where the text is "
        "not readable, so they are worth the most."
    )

    for i, (page_no, block, ratio) in enumerate(damaged, start=1):
        img = os.path.join(OUT, f"garbled-{page_no}-{i}.jpg")
        try:
            block_strip(scan, pages, page_no, block, img)
        except Exception as exc:
            print(f"  skipped passage on page {page_no}: {exc}")
            continue
        doc.add_paragraph()
        head_p = doc.add_paragraph()
        head_p.add_run(f"{i}. ").bold = True
        head_p.add_run(f"scan page {page_no}").bold = True
        head_p.add_run(f"   ·   {ratio:.0%} of this passage is unreadable")
        doc.add_picture(img, width=Inches(6.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cur = doc.add_paragraph()
        cur.add_run("Site now:  ").bold = True
        shade(cur.add_run(re.sub(r"\s+", " ", block)[:420]), (0x44, 0x44, 0x44))
        ask = doc.add_paragraph()
        ask.add_run("On the page:  ").bold = True
        ask.add_run("_" * 74)
        doc.add_paragraph("_" * 92)
        doc.add_paragraph("_" * 92)

    # ---- part two: single words that could be read either way -------------
    doc.add_page_break()
    doc.add_heading("Part two — single words", 1)
    doc.add_paragraph(
        f"{len(batch)} of {len(suspects)} words that are real words in their "
        "own right, so only the page can settle them. Most are probably right "
        "as they stand."
    )

    made = 0
    for n, (word, freq, near, _) in enumerate(batch, start=1):
        found = find_word(pages, word)
        if not found:
            continue
        page_no, line_no, line_count, _raw = found
        img = os.path.join(OUT, f"p{page_no}-{word}.jpg")
        try:
            strip_for(scan, page_no, line_no, line_count, img)
        except Exception as exc:
            print(f"  skipped {word}: {exc}")
            continue

        doc.add_paragraph()
        head_p = doc.add_paragraph()
        head_p.add_run(f"{n}. ").bold = True
        r = head_p.add_run(f'"{word}"')
        r.bold = True
        r.font.size = Pt(14)
        head_p.add_run(f"   ·   scan page {page_no}")

        doc.add_picture(img, width=Inches(6.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cur = doc.add_paragraph()
        cur.add_run("Site now:  ").bold = True
        shade(cur.add_run(sentence_for(word, texts)), (0x44, 0x44, 0x44))

        alt = doc.add_paragraph()
        alt.add_run("Or could be:  ").bold = True
        shade(alt.add_run(", ".join(near)), (0x44, 0x44, 0x44))

        ask = doc.add_paragraph()
        ask.add_run("On the page:  ").bold = True
        ask.add_run("_" * 74)
        made += 1

    path = os.path.join(OUT, "wastena-review.docx")
    doc.save(path)
    print(f"wrote {path}")
    print(f"  part one: {len(damaged)} damaged passages")
    print(f"  part two: {made} single words")


if __name__ == "__main__":
    main()
